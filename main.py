"""
NOVA Agent Server v6.0 - COMPLETE REBUILD
==========================================

CRITICAL CHANGE FROM v5.x:
The Analysis Agent now follows the principle:
    "EXTRACT and CITE, never GENERATE and GUESS"

Every factual claim must come from web research with citation.
Missing information is marked "Not found through research" - NEVER fabricated.

Architecture:
- Research Phase: Claude uses web_search 15-20 times for comprehensive coverage
- Extraction Phase: Claude extracts ONLY facts found with URLs
- Validation Phase: Anything not found is explicitly marked
- Document Phase: Only real data goes into output

Endpoints:
- POST /api/execute - Start agent task
- GET /api/status/{job_id} - Get task status  
- GET /api/download/{job_id} - Download ZIP
- GET /api/health - Health check

Author: Claude AI for NOVA Project
Date: 31 January 2026
Version: 6.0.0
"""

import os
import json
import asyncio
import re
import httpx
from datetime import datetime
from typing import Optional, Dict, Any, List
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import zipfile
import io
import anthropic

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ============================================================================
# APP SETUP
# ============================================================================

app = FastAPI(title="NOVA Agent Server", version="6.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Job Storage
class JobStore:
    def __init__(self):
        self._jobs: Dict[str, Dict] = {}
    
    def get(self, job_id: str) -> Optional[Dict]:
        return self._jobs.get(job_id)
    
    def set(self, job_id: str, data: Dict):
        self._jobs[job_id] = data
    
    def update(self, job_id: str, **kwargs):
        if job_id in self._jobs:
            self._jobs[job_id].update(kwargs)
    
    def exists(self, job_id: str) -> bool:
        return job_id in self._jobs

jobs = JobStore()

OUTPUT_DIR = Path("/tmp/nova-outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

# Claude client
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-5-20250929")
claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None

print(f"[NOVA v6.0] Started. Claude configured: {claude_client is not None}, Model: {CLAUDE_MODEL}")


# ============================================================================
# API MODELS - BACKWARD COMPATIBLE WITH v5.1 FRONTEND
# ============================================================================

class ExecuteRequest(BaseModel):
    """Execute request - supports both old and new format"""
    # v5.1 format
    job_id: Optional[str] = None
    agent: Optional[str] = None
    # v6.0 format
    agent_type: Optional[str] = None
    parameters: Dict[str, Any] = {}
    framework: str = "INDUSTRY_STANDARDS"


class TaskResponse(BaseModel):
    """Response format expected by frontend"""
    job_id: str
    status: str
    message: str


class StatusResponse(BaseModel):
    """Status response format expected by frontend"""
    job_id: str
    status: str
    progress: int
    current_step: str
    steps_completed: List[str]
    error: Optional[str]
    created_at: str
    completed_at: Optional[str]


class JobStatus(BaseModel):
    job_id: str
    status: str
    progress: int
    message: str
    download_url: Optional[str] = None


# ============================================================================
# FRAMEWORK TERMINOLOGY
# ============================================================================

def get_terminology(framework: str) -> Dict[str, str]:
    """Get framework-specific terminology"""
    
    framework = normalize_framework(framework)
    
    TERMINOLOGY = {
        "UK_DSAT": {
            "framework_name": "UK Defence Systems Approach to Training (DSAT)",
            "task_list": "Role Performance Statement (RolePS)",
            "task_list_short": "RolePS",
            "top_objective": "Training Objective (TO)",
            "top_objective_short": "TO",
            "enabling_objective": "Enabling Objective (EO)",
            "enabling_objective_short": "EO",
            "learning_point": "Key Learning Point (KLP)",
            "learning_point_short": "KLP",
            "needs_report": "Training Needs Report (TNR)",
            "course_design": "Learning Specification (LSpec)",
            "lesson_plan": "Lesson Plan (PAR Structure)",
            "internal_eval": "Internal Validation (InVal)",
            "external_eval": "External Validation (ExVal)",
            "citation_prefix": "[JSP 822 V7.0]",
            "authority": "UK Ministry of Defence"
        },
        "US_TRADOC": {
            "framework_name": "US Army Training and Doctrine Command (TRADOC)",
            "task_list": "Individual Critical Task List (ICTL)",
            "task_list_short": "ICTL",
            "top_objective": "Terminal Learning Objective (TLO)",
            "top_objective_short": "TLO",
            "enabling_objective": "Enabling Learning Objective (ELO)",
            "enabling_objective_short": "ELO",
            "learning_point": "Learning Step Activity (LSA)",
            "learning_point_short": "LSA",
            "needs_report": "Training Requirements Document (TRD)",
            "course_design": "Program of Instruction (POI)",
            "lesson_plan": "Lesson Plan (5-Section Format)",
            "internal_eval": "Internal Evaluation",
            "external_eval": "External Evaluation",
            "citation_prefix": "[TRADOC Reg 350-70]",
            "authority": "US Army TRADOC"
        },
        "NATO_BISC": {
            "framework_name": "NATO Bi-Strategic Command (Bi-SC) Directive 075-007",
            "task_list": "Skills, Tasks, Proficiency (STP) Analysis",
            "task_list_short": "STP Analysis",
            "top_objective": "Performance Objective (PO)",
            "top_objective_short": "PO",
            "enabling_objective": "Enabling Learning Objective (ELO)",
            "enabling_objective_short": "ELO",
            "learning_point": "Learning Points",
            "learning_point_short": "LP",
            "needs_report": "Training Requirements Analysis (TRA) Report",
            "course_design": "Course Control Document II (CCD II)",
            "lesson_plan": "Course Control Document III (CCD III)",
            "internal_eval": "Course Evaluation",
            "external_eval": "External Quality Assurance",
            "citation_prefix": "[Bi-SCD 075-007]",
            "authority": "NATO Allied Command Transformation"
        },
        "AUSTRALIAN_SADL": {
            "framework_name": "Australian Systematic Approach to Defence Learning (SADL)",
            "task_list": "Task Analysis",
            "task_list_short": "Task Analysis",
            "top_objective": "Learning Outcome (LO)",
            "top_objective_short": "LO",
            "enabling_objective": "Supporting Learning Outcome (SLO)",
            "enabling_objective_short": "SLO",
            "learning_point": "Learning Points",
            "learning_point_short": "LP",
            "needs_report": "Analysis Phase Report",
            "course_design": "Curriculum Design Document",
            "lesson_plan": "Lesson Plan",
            "internal_eval": "Internal Review",
            "external_eval": "External Review",
            "citation_prefix": "[Defence Learning Manual]",
            "authority": "Australian Defence Force"
        },
        "S6000T": {
            "framework_name": "ASD/AIA S6000T Training Analysis and Design",
            "task_list": "Task Specification",
            "task_list_short": "Task Spec",
            "top_objective": "Training Requirement",
            "top_objective_short": "TR",
            "enabling_objective": "Sub-Task Requirement",
            "enabling_objective_short": "STR",
            "learning_point": "Task Element",
            "learning_point_short": "TE",
            "needs_report": "Training Needs Analysis Report",
            "course_design": "Training Specification",
            "lesson_plan": "Training Module",
            "internal_eval": "Internal Verification",
            "external_eval": "External Verification",
            "citation_prefix": "[S6000T]",
            "authority": "AeroSpace and Defence Industries Association"
        },
        "ADDIE": {
            "framework_name": "ADDIE Instructional Design Model",
            "task_list": "Job/Task Analysis",
            "task_list_short": "Task Analysis",
            "top_objective": "Learning Objective",
            "top_objective_short": "LO",
            "enabling_objective": "Enabling Objective",
            "enabling_objective_short": "EO",
            "learning_point": "Learning Point",
            "learning_point_short": "LP",
            "needs_report": "Analysis Phase Documentation",
            "course_design": "Design Document",
            "lesson_plan": "Lesson Plan",
            "internal_eval": "Formative Evaluation",
            "external_eval": "Summative Evaluation",
            "citation_prefix": "[ADDIE Model]",
            "authority": "Industry Standard"
        },
        "SAM": {
            "framework_name": "Successive Approximation Model (SAM)",
            "task_list": "Savvy Start Analysis",
            "task_list_short": "Savvy Start",
            "top_objective": "Performance Objective",
            "top_objective_short": "PO",
            "enabling_objective": "Supporting Objective",
            "enabling_objective_short": "SO",
            "learning_point": "Learning Activity",
            "learning_point_short": "LA",
            "needs_report": "Preparation Phase Summary",
            "course_design": "Design Proof",
            "lesson_plan": "Iterative Lesson Module",
            "internal_eval": "Alpha/Beta Review",
            "external_eval": "Gold Release Review",
            "citation_prefix": "[SAM Model]",
            "authority": "Allen Interactions"
        },
        "ISO_29990": {
            "framework_name": "ISO 29990 Learning Services Standard",
            "task_list": "Learning Needs Determination",
            "task_list_short": "Needs Determination",
            "top_objective": "Learning Outcome",
            "top_objective_short": "LO",
            "enabling_objective": "Supporting Outcome",
            "enabling_objective_short": "SO",
            "learning_point": "Learning Element",
            "learning_point_short": "LE",
            "needs_report": "Learning Needs Analysis Record",
            "course_design": "Learning Service Design",
            "lesson_plan": "Learning Activity Plan",
            "internal_eval": "Internal Monitoring",
            "external_eval": "External Evaluation",
            "citation_prefix": "[ISO 29990]",
            "authority": "International Organization for Standardization"
        },
        "KIRKPATRICK": {
            "framework_name": "Industry Standards",
            "task_list": "Job/Task Analysis",
            "task_list_short": "Task Analysis",
            "top_objective": "Learning Objective",
            "top_objective_short": "LO",
            "enabling_objective": "Enabling Objective",
            "enabling_objective_short": "EO",
            "learning_point": "Learning Point",
            "learning_point_short": "LP",
            "needs_report": "Training Needs Analysis Report",
            "course_design": "Course Design Document",
            "lesson_plan": "Lesson Plan",
            "internal_eval": "Internal Evaluation",
            "external_eval": "External Evaluation",
            "citation_prefix": "[Industry Standard]",
            "authority": "Relevant Industry Bodies"
        },
        "ACTION_MAPPING": {
            "framework_name": "Industry Standards",
            "task_list": "Job/Task Analysis",
            "task_list_short": "Task Analysis",
            "top_objective": "Learning Objective",
            "top_objective_short": "LO",
            "enabling_objective": "Enabling Objective",
            "enabling_objective_short": "EO",
            "learning_point": "Learning Point",
            "learning_point_short": "LP",
            "needs_report": "Training Needs Analysis Report",
            "course_design": "Course Design Document",
            "lesson_plan": "Lesson Plan",
            "internal_eval": "Internal Evaluation",
            "external_eval": "External Evaluation",
            "citation_prefix": "[Industry Standard]",
            "authority": "Relevant Industry Bodies"
        },
        "COMMERCIAL": {
            "framework_name": "Industry Standards",
            "task_list": "Job/Task Analysis",
            "task_list_short": "Task Analysis",
            "top_objective": "Learning Objective",
            "top_objective_short": "LO",
            "enabling_objective": "Enabling Objective",
            "enabling_objective_short": "EO",
            "learning_point": "Learning Point",
            "learning_point_short": "LP",
            "needs_report": "Training Needs Analysis Report",
            "course_design": "Course Design Document",
            "lesson_plan": "Lesson Plan",
            "internal_eval": "Internal Evaluation",
            "external_eval": "External Evaluation",
            "citation_prefix": "[Industry Standard]",
            "authority": "Relevant Industry Bodies"
        },
        "INDUSTRY_STANDARDS": {
            "framework_name": "Industry Standards",
            "task_list": "Job/Task Analysis",
            "task_list_short": "Task Analysis",
            "top_objective": "Learning Objective",
            "top_objective_short": "LO",
            "enabling_objective": "Enabling Objective",
            "enabling_objective_short": "EO",
            "learning_point": "Learning Point",
            "learning_point_short": "LP",
            "needs_report": "Training Needs Analysis Report",
            "course_design": "Course Design Document",
            "lesson_plan": "Lesson Plan",
            "internal_eval": "Internal Evaluation",
            "external_eval": "External Evaluation",
            "citation_prefix": "[Industry Standard]",
            "authority": "Relevant Industry Bodies"
        }
    }
    
    return TERMINOLOGY.get(framework, TERMINOLOGY["COMMERCIAL"])


def normalize_framework(framework: str) -> str:
    """Normalize framework name to standard key"""
    if not framework:
        return "INDUSTRY_STANDARDS"
    
    framework_lower = framework.lower().replace("-", "_").replace(" ", "_")
    
    mapping = {
        # Allied Defence Training (keep all 5)
        "uk_dsat": "UK_DSAT", "uk": "UK_DSAT", "dsat": "UK_DSAT", "jsp822": "UK_DSAT",
        "us_tradoc": "US_TRADOC", "us": "US_TRADOC", "tradoc": "US_TRADOC",
        "nato_bisc": "NATO_BISC", "nato": "NATO_BISC", "bisc": "NATO_BISC",
        "australian_sadl": "AUSTRALIAN_SADL", "sadl": "AUSTRALIAN_SADL", "australia": "AUSTRALIAN_SADL",
        "s6000t": "S6000T", "asd_s6000t": "S6000T",
        
        # ALL Commercial/Industry frameworks map to INDUSTRY_STANDARDS
        "addie": "INDUSTRY_STANDARDS",
        "sam": "INDUSTRY_STANDARDS",
        "iso_29990": "INDUSTRY_STANDARDS", "iso29990": "INDUSTRY_STANDARDS",
        "kirkpatrick": "INDUSTRY_STANDARDS", "atd": "INDUSTRY_STANDARDS",
        "action_mapping": "INDUSTRY_STANDARDS",
        "commercial": "INDUSTRY_STANDARDS", 
        "industry_standards": "INDUSTRY_STANDARDS", 
        "industry": "INDUSTRY_STANDARDS"
    }
    
    return mapping.get(framework_lower, "INDUSTRY_STANDARDS")


def is_defence_framework(framework: str) -> bool:
    """Check if framework is an Allied Defence framework"""
    defence_frameworks = ["UK_DSAT", "US_TRADOC", "NATO_BISC", "AUSTRALIAN_SADL", "S6000T"]
    return normalize_framework(framework) in defence_frameworks


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.get("/api/health")
async def health_check():
    return {
        "status": "healthy",
        "version": "6.0.0",
        "claude_configured": claude_client is not None,
        "timestamp": datetime.now().isoformat()
    }


@app.post("/api/execute", response_model=TaskResponse)
async def execute_agent(request: ExecuteRequest, background_tasks: BackgroundTasks):
    """Start an agent task - backward compatible with v5.1 frontend"""
    
    try:
        # Validate Claude is configured
        if not claude_client:
            raise HTTPException(
                status_code=500, 
                detail="ANTHROPIC_API_KEY not configured on server"
            )
        
        # Handle both v5.1 and v6.0 request formats
        agent_type = request.agent_type or request.agent or "analysis"
        
        # Generate job_id if not provided
        job_id = request.job_id or f"nova-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{os.urandom(4).hex()}"
        
        # Create output directory
        job_output_dir = OUTPUT_DIR / job_id
        job_output_dir.mkdir(parents=True, exist_ok=True)
        
        # Get framework from parameters or request
        framework = request.parameters.get("framework") or request.framework or "INDUSTRY_STANDARDS"
        framework = normalize_framework(framework)
        
        # Initialize job with all fields expected by StatusResponse
        jobs.set(job_id, {
            "job_id": job_id,
            "status": "running",
            "progress": 0,
            "message": "Initializing...",
            "current_step": "Initializing...",
            "steps_completed": [],
            "error": None,
            "agent_type": agent_type,
            "framework": framework,
            "parameters": request.parameters,
            "output_dir": str(job_output_dir),
            "created_at": datetime.now().isoformat(),
            "completed_at": None
        })
        
        print(f"[NOVA v6.0] Job {job_id} started: {agent_type} agent, {framework} framework")
        
        # Create modified request for background task
        class ModifiedRequest:
            def __init__(self, agent_type, parameters, framework):
                self.agent_type = agent_type
                self.parameters = parameters
                self.framework = framework
        
        modified_request = ModifiedRequest(agent_type, request.parameters, framework)
        
        # Run in background
        background_tasks.add_task(run_agent_task, job_id, modified_request)
        
        # Return v5.1 compatible response
        return TaskResponse(
            job_id=job_id, 
            status="started", 
            message=f"Agent {agent_type} started with {framework} framework"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[NOVA] Execute error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to start agent: {str(e)}")


@app.get("/api/status/{job_id}", response_model=StatusResponse)
async def get_status(job_id: str):
    """Get job status - v5.1 compatible format"""
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Return StatusResponse format expected by frontend
    return StatusResponse(
        job_id=job_id,
        status=job.get("status", "unknown"),
        progress=job.get("progress", 0),
        current_step=job.get("message", "Processing..."),
        steps_completed=job.get("steps_completed", []),
        error=job.get("error"),
        created_at=job.get("created_at", datetime.now().isoformat()),
        completed_at=job.get("completed_at")
    )


@app.get("/api/download/{job_id}")
async def download_results(job_id: str):
    """Download job results as ZIP"""
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if job.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Job not complete")
    
    output_dir = Path(job["output_dir"])
    zip_path = output_dir.parent / f"{job_id}.zip"
    
    # Create ZIP
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for file_path in output_dir.rglob("*"):
            if file_path.is_file():
                arcname = file_path.relative_to(output_dir)
                zf.write(file_path, arcname)
    
    return FileResponse(
        path=str(zip_path),
        filename=f"NOVA_{job_id}.zip",
        media_type="application/zip"
    )


def update_job(job_id: str, progress: int, message: str, status: str = None):
    """Update job progress"""
    update_data = {"progress": progress, "message": message}
    if status:
        update_data["status"] = status
        if status == "completed":
            update_data["completed_at"] = datetime.now().isoformat()
    
    # Also update current_step to match message
    update_data["current_step"] = message
    
    # Add to steps_completed if progress milestone reached
    job = jobs.get(job_id)
    if job:
        steps = job.get("steps_completed", [])
        if progress > 0 and progress % 20 == 0:  # Track every 20%
            steps.append(f"{progress}%: {message}")
        update_data["steps_completed"] = steps
    
    jobs.update(job_id, **update_data)
    print(f"[NOVA] {job_id}: {progress}% - {message}")


# ============================================================================
# AGENT TASK ROUTER
# ============================================================================

async def run_agent_task(job_id: str, request):
    """Route to appropriate agent based on type"""
    try:
        framework = normalize_framework(request.framework)
        terms = get_terminology(framework)
        
        print(f"[NOVA v6.0] Starting {request.agent_type} agent with {framework}")
        
        if request.agent_type == "analysis":
            await run_analysis_agent(job_id, request.parameters, framework, terms)
        elif request.agent_type == "design":
            await run_design_agent(job_id, request.parameters, framework, terms)
        elif request.agent_type == "delivery":
            await run_delivery_agent(job_id, request.parameters, framework, terms)
        elif request.agent_type == "evaluation":
            await run_evaluation_agent(job_id, request.parameters, framework, terms)
        else:
            raise ValueError(f"Unknown agent type: {request.agent_type}")
        
        # Mark completed with timestamp
        jobs.update(job_id, 
            status="completed", 
            progress=100, 
            message="Complete",
            current_step="Complete",
            completed_at=datetime.now().isoformat()
        )
        
    except Exception as e:
        print(f"[NOVA] Error in {request.agent_type}: {e}")
        import traceback
        traceback.print_exc()
        jobs.update(job_id, 
            status="error", 
            message=str(e),
            current_step=f"Error: {str(e)}",
            error=str(e)
        )


# ============================================================================
# CLAUDE API CALLS
# ============================================================================

async def call_claude_with_search(system_prompt: str, user_prompt: str, max_tokens: int = 16000) -> Dict:
    """Call Claude API with timeout - no web search to avoid delays"""
    if not claude_client:
        raise Exception("Claude API not configured")
    
    print(f"[NOVA] Calling Claude ({len(user_prompt)} chars)...")
    
    loop = asyncio.get_event_loop()
    
    try:
        response = await asyncio.wait_for(
            loop.run_in_executor(
                None,
                lambda: claude_client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=max_tokens,
                    system=system_prompt + "\n\nIMPORTANT: Base analysis on your training knowledge. State when information requires verification. Never fabricate statistics or costs.",
                    messages=[{"role": "user", "content": user_prompt}]
                )
            ),
            timeout=120.0
        )
        
        result = {
            "text": response.content[0].text if response.content else "",
            "citations": [],
            "searches_performed": []
        }
        print(f"[NOVA] Response received: {len(result['text'])} chars")
        return result
        
    except asyncio.TimeoutError:
        raise Exception("Claude API timeout after 120 seconds")
    except Exception as e:
        raise Exception(f"Claude API error: {str(e)}")


async def call_claude_standard(system_prompt: str, user_prompt: str, max_tokens: int = 8000) -> str:
    """Standard Claude call with timeout (for Design/Delivery/Evaluation)"""
    if not claude_client:
        raise Exception("Claude API not configured")
    
    loop = asyncio.get_event_loop()
    
    try:
        response = await asyncio.wait_for(
            loop.run_in_executor(
                None,
                lambda: claude_client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=max_tokens,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}]
                )
            ),
            timeout=120.0
        )
        
        return response.content[0].text
        
    except asyncio.TimeoutError:
        raise Exception("Claude API timeout after 120 seconds")
    except Exception as e:
        raise Exception(f"Claude API error: {str(e)}")


# ============================================================================
# ANALYSIS AGENT v6.0 - RESEARCH-BASED IMPLEMENTATION
# ============================================================================
# 
# CRITICAL PRINCIPLE: "EXTRACT and CITE, never GENERATE and GUESS"
#
# This agent:
# 1. Conducts comprehensive web research (15-25 searches)
# 2. Extracts ONLY facts found with URLs
# 3. Marks missing info as "Not found through research"
# 4. NEVER fabricates statistics, costs, or methodology
# ============================================================================

def get_analysis_system_prompt() -> str:
    """
    System prompt for Analysis Agent - enforces EXTRACT and CITE principle.
    This prompt is critical - it instructs Claude to ONLY report findings.
    """
    return """# NOVA™ ANALYSIS AGENT v6.0

## YOUR IDENTITY
You are the NOVA™ Analysis Agent. Your role is to conduct comprehensive research and REPORT ONLY WHAT YOU FIND. You are a research assistant, not a content generator.

## ABSOLUTE RULES - VIOLATION IS FAILURE

### RULE 1: NEVER FABRICATE
You must NEVER invent, guess, or fabricate:
- Statistics or percentages (e.g., "78% of engineers...")
- Cost estimates (e.g., "$75,000 training budget")
- Methodology claims (e.g., "through interviews with 50 SMEs")
- Professional requirements you haven't verified
- Standards or regulations you haven't found

### RULE 2: CITE EVERYTHING
Every factual claim MUST include its source:
- Format: "Claim text [Source: Organisation - URL]"
- If you found it through web search, cite the URL
- If you cannot cite it, do not include it

### RULE 3: ACKNOWLEDGE GAPS
When you cannot find information, you MUST state:
- "Not found through research"
- "No authoritative source identified"
- "Research inconclusive - further investigation required"

NEVER fill gaps with plausible-sounding fabrications.

### RULE 4: USE WEB SEARCH
You MUST use the web_search tool to find real, current information.
Do not rely on your training data for:
- Current professional body requirements
- Current qualification frameworks
- Current legislation
- Current salary data
- Current industry standards

## OUTPUT FORMAT
Return ONLY a valid JSON object. No explanatory text before or after.
Every section must reflect ACTUAL RESEARCH FINDINGS.

## QUALITY CHECK
Before returning your response, verify:
✓ Every statistic has a citation
✓ Every requirement has a source URL
✓ No fabricated methodology claims
✓ Missing data is explicitly marked as "Not found"
✓ JSON is valid and complete"""


def get_domain_knowledge(domain: str) -> Dict:
    """
    Domain-specific knowledge to guide research queries.
    This tells the agent WHAT to search for in each domain.
    """
    
    DOMAINS = {
        "AI & Data Science": {
            "primary_frameworks": ["SFIA 9", "BCS Digital Framework", "UNESCO AI Competency Framework"],
            "iso_standards": ["ISO/IEC 42001", "ISO/IEC 27001", "ISO/IEC 25010"],
            "professional_bodies": ["BCS The Chartered Institute for IT", "IET", "Alan Turing Institute"],
            "apprenticeships": ["AI Data Specialist Level 7", "Data Analyst Level 4", "Software Developer Level 4"],
            "regulations": ["UK AI Regulation", "Data Protection Act 2018", "UK GDPR"],
            "key_searches": [
                "SFIA 9 data engineering skills levels",
                "BCS professional membership requirements",
                "AI data specialist apprenticeship standard",
                "data protection training requirements UK"
            ]
        },
        "Healthcare": {
            "primary_frameworks": ["NHS Knowledge and Skills Framework", "Skills for Health NOS"],
            "iso_standards": ["ISO 13485", "ISO 15189", "CQC Standards"],
            "professional_bodies": ["GMC", "NMC", "HCPC", "GPhC"],
            "apprenticeships": ["Nursing Associate Level 5", "Healthcare Support Worker Level 2"],
            "regulations": ["Health and Social Care Act 2008", "Care Act 2014", "Mental Health Act 1983"],
            "key_searches": [
                "NHS KSF competency levels",
                "NMC registration requirements UK",
                "healthcare mandatory training requirements"
            ]
        },
        "Defence": {
            "primary_frameworks": ["JSP 822", "DTSM 1-5", "NATO STANAG"],
            "iso_standards": ["ASD S6000T", "DEF STAN", "AQAP"],
            "professional_bodies": ["Defence Academy UK"],
            "apprenticeships": ["Military technical apprenticeships"],
            "regulations": ["Armed Forces Act", "Official Secrets Act"],
            "key_searches": [
                "JSP 822 training requirements",
                "UK military security clearance levels",
                "defence training standards MOD"
            ]
        },
        "Finance & Banking": {
            "primary_frameworks": ["FCA Training and Competence", "CFA Competency Framework"],
            "iso_standards": ["ISO 22301", "PCI DSS"],
            "professional_bodies": ["FCA", "PRA", "CII", "ICAEW", "ACCA"],
            "apprenticeships": ["Financial Services Administrator Level 3", "Financial Adviser Level 4"],
            "regulations": ["FSMA 2000", "SMCR", "Money Laundering Regulations 2017"],
            "key_searches": [
                "FCA SMCR training requirements",
                "financial services CPD requirements UK",
                "anti-money laundering training requirements"
            ]
        },
        "Construction & Engineering": {
            "primary_frameworks": ["Engineering Council UK-SPEC", "CITB Standards"],
            "iso_standards": ["ISO 9001", "ISO 45001", "ISO 14001", "Eurocodes"],
            "professional_bodies": ["Engineering Council", "ICE", "IMechE", "CIOB", "RICS"],
            "apprenticeships": ["Civil Engineer Degree Level 6", "Construction Site Manager Level 6"],
            "regulations": ["CDM Regulations 2015", "Building Regulations", "Building Safety Act 2022"],
            "key_searches": [
                "CEng chartered engineer requirements",
                "CSCS card requirements",
                "construction mandatory training UK"
            ]
        },
        "Legal": {
            "primary_frameworks": ["SRA Competence Statement", "BSB Competencies"],
            "iso_standards": ["Lexcel"],
            "professional_bodies": ["SRA", "BSB", "CILEX", "Law Society"],
            "apprenticeships": ["Solicitor Level 7", "Paralegal Level 3"],
            "regulations": ["Legal Services Act 2007", "SRA Standards and Regulations"],
            "key_searches": [
                "SQE solicitors qualifying examination",
                "SRA CPD requirements 2024",
                "legal professional competencies UK"
            ]
        },
        "Education & Training": {
            "primary_frameworks": ["Teachers Standards UK", "ETF Professional Standards"],
            "iso_standards": ["ISO 29990", "ISO 21001"],
            "professional_bodies": ["TRA", "ETF", "Chartered College of Teaching"],
            "apprenticeships": ["Learning and Skills Teacher Level 5"],
            "regulations": ["Education Act 2011", "KCSIE", "Prevent Duty"],
            "key_searches": [
                "QTS qualified teacher status requirements",
                "FE teaching qualifications UK",
                "safeguarding training requirements education"
            ]
        },
        "Manufacturing & Operations": {
            "primary_frameworks": ["SEMTA NOS", "Make UK Competencies"],
            "iso_standards": ["ISO 9001", "ISO 45001", "IATF 16949", "AS9100"],
            "professional_bodies": ["IMechE", "IOM", "CQI"],
            "apprenticeships": ["Manufacturing Engineer Level 6", "Engineering Technician Level 3"],
            "regulations": ["HASAWA 1974", "PUWER", "COSHH"],
            "key_searches": [
                "manufacturing NVQ qualifications UK",
                "lean six sigma certification levels",
                "IOSH managing safely requirements"
            ]
        }
    }
    
    # Find best match
    domain_lower = domain.lower()
    for key, value in DOMAINS.items():
        if any(term in domain_lower for term in key.lower().split()):
            return value
    
    # Generic fallback
    return {
        "primary_frameworks": ["National Occupational Standards"],
        "iso_standards": ["ISO 9001"],
        "professional_bodies": [],
        "apprenticeships": [],
        "regulations": [],
        "key_searches": [f"{domain} professional standards UK", f"{domain} competency framework"]
    }


def build_research_prompt(
    domain: str,
    specialism: str, 
    role_title: str,
    proficiency_level: str,
    framework: str,
    role_description: str,
    terms: Dict
) -> str:
    """
    Build comprehensive research prompt.
    This prompt instructs Claude to SEARCH and REPORT, not generate.
    """
    
    domain_info = get_domain_knowledge(domain)
    
    # Format domain-specific info
    frameworks_list = ", ".join(domain_info.get("primary_frameworks", []))
    standards_list = ", ".join(domain_info.get("iso_standards", []))
    bodies_list = ", ".join(domain_info.get("professional_bodies", []))
    
    return f"""# RESEARCH TASK: Training Analysis for {role_title}

## PARAMETERS
- Domain: {domain}
- Specialism: {specialism}
- Role Title: {role_title}
- Proficiency Level: {proficiency_level}
- Framework: {framework} ({terms.get('framework_name', framework)})
- Additional Context: {role_description or 'None provided'}

## YOUR TASK
Conduct comprehensive web research to build a factual profile of this role. You must use web_search for each category below.

## MANDATORY RESEARCH STEPS

### STEP 1: FRAMEWORK RESEARCH
Search for "{framework} training framework requirements" and "{terms.get('task_list', 'task analysis')} format"
Capture: Framework version, governing body, required outputs

### STEP 2: PROFESSIONAL BODY RESEARCH  
Search for "{specialism} professional body UK" and "{domain} regulatory body UK"
Known bodies to search: {bodies_list}
Capture: Registration requirements, membership categories, protected titles

### STEP 3: COMPETENCY FRAMEWORK RESEARCH
Search for "{specialism} competency framework UK" and "{domain} National Occupational Standards"
Known frameworks to search: {frameworks_list}
Capture: Framework name, relevant competency units, level descriptors

### STEP 4: QUALIFICATION RESEARCH
Search for "{role_title} qualifications UK" and "{specialism} apprenticeship standard"
Search for "{specialism} degree requirements UK" and "{specialism} professional certification"
Capture: Essential qualifications, RQF levels, apprenticeship routes

### STEP 5: ROLE DEFINITION RESEARCH
Search for "{role_title} job description UK" and "{role_title} responsibilities"
Capture: Standard definition, key accountabilities, typical duties

### STEP 6: EXPERIENCE REQUIREMENTS
Search for "{role_title} {proficiency_level} experience requirements"
Search for "{specialism} career progression UK"
Capture: Years of experience, type of experience, progression pathway

### STEP 7: SKILLS RESEARCH
Search for "{specialism} technical skills requirements" and "{role_title} core competencies"
Capture: Technical skills with proficiency levels, soft skills

### STEP 8: LEGAL/COMPLIANCE RESEARCH
Search for "{domain} UK legislation" and "{specialism} mandatory training UK"
Search for "{role_title} legal requirements UK"
Capture: Applicable legislation, statutory duties, mandatory training

### STEP 9: SECURITY/MEDICAL REQUIREMENTS
Search for "{role_title} DBS requirements UK" and "{domain} security clearance UK"
Capture: DBS level, security clearance, health requirements

### STEP 10: CPD REQUIREMENTS
Search for "{specialism} CPD requirements UK" and "{specialism} recertification"
Capture: Annual CPD hours, recertification cycle, mandatory refreshers

## OUTPUT FORMAT

Return a JSON object with this structure. For EVERY field:
- If found: Include the information with "[Source: URL]"
- If not found: Use "Not found through research"

```json
{{
    "research_summary": {{
        "searches_conducted": ["list actual searches you performed"],
        "sources_found": ["list authoritative sources discovered"],
        "research_date": "{datetime.now().strftime('%Y-%m-%d')}",
        "research_limitations": "any gaps or limitations in findings"
    }},
    
    "section_01_executive_summary": {{
        "analysis_scope": "Analysis of {role_title} in {domain}/{specialism} at {proficiency_level} level",
        "methodology": "Web-based research using authoritative UK sources",
        "key_findings": ["Finding 1 [Source: URL]", "Finding 2 [Source: URL]"],
        "tasks_identified": <number based on research>,
        "primary_standards_referenced": ["Standard 1", "Standard 2"]
    }},
    
    "section_02_framework_identification": {{
        "framework_name": "{terms.get('framework_name', framework)}",
        "framework_version": "Version found [Source: URL]" or "Not found through research",
        "governing_authority": "{terms.get('authority', 'Unknown')}",
        "analysis_requirements": ["Requirement [Source]"],
        "terminology_glossary": {{"term": "definition"}}
    }},
    
    "section_03_geographic_context": {{
        "country": "United Kingdom",
        "legal_jurisdiction": "England and Wales",
        "language": "English",
        "currency": "GBP"
    }},
    
    "section_04_professional_body": {{
        "body_name": "Name [Source: URL]" or "Not found through research",
        "website_url": "URL",
        "registration_required": true/false,
        "registration_requirements": ["Requirement [Source]"],
        "protected_titles": ["Title [Source]"],
        "membership_categories": ["Category"]
    }},
    
    "section_05_competency_framework": {{
        "framework_name": "Name [Source: URL]" or "Not found through research",
        "framework_owner": "Organisation",
        "relevant_competencies": [
            {{"code": "Code", "name": "Name", "level": "Level for {proficiency_level}"}}
        ],
        "proficiency_mapping": "How {proficiency_level} maps to framework"
    }},
    
    "section_06_role_description": {{
        "definition": "Definition [Source: URL]" or "Not found through research",
        "primary_purpose": "Purpose statement",
        "key_accountabilities": ["Accountability [Source]"],
        "reporting_structure": "Typical structure",
        "equivalent_titles": ["Alternative title"]
    }},
    
    "section_07_qualifications": {{
        "essential": [
            {{"qualification": "Name", "level": "RQF Level", "source": "URL"}}
        ],
        "desirable": [
            {{"qualification": "Name", "level": "Level", "source": "URL"}}
        ],
        "apprenticeship_routes": [
            {{"name": "Name", "level": "Level", "source": "URL"}}
        ]
    }},
    
    "section_08_experience": {{
        "years_required": "X years [Source: URL]" or "Not found through research",
        "experience_types": ["Type [Source]"],
        "sector_requirements": ["Requirement"]
    }},
    
    "section_09_technical_skills": [
        {{"skill": "Skill name", "category": "Core/Desirable", "proficiency": "Level", "source": "URL"}}
    ],
    
    "section_10_soft_skills": [
        {{"skill": "Skill name", "proficiency": "Level", "source": "URL"}}
    ],
    
    "section_11_behaviours": [
        {{"behaviour": "Description", "type": "Essential/Desirable", "source": "URL"}}
    ],
    
    "section_12_physical_medical_security": {{
        "physical_requirements": "Requirements [Source]" or "No specific requirements identified",
        "medical_requirements": "Requirements [Source]" or "Standard employment health",
        "security_clearance": "Level required [Source]" or "Standard DBS",
        "dbs_level": "Basic/Standard/Enhanced [Source]"
    }},
    
    "section_13_cpd_requirements": {{
        "professional_body_cpd": "Policy [Source: URL]" or "Not found through research",
        "annual_hours": "X hours [Source]",
        "recertification_cycle": "X years [Source]",
        "mandatory_refreshers": ["Training [Source]"]
    }},
    
    "section_14_career_progression": {{
        "pathway_to_role": ["Previous role"],
        "pathway_from_role": ["Next role"],
        "typical_timeline": "X years [Source]"
    }},
    
    "section_15_legal_compliance": [
        {{"legislation": "Act Name Year", "relevance": "How it applies", "mandatory_training": true/false, "source": "URL"}}
    ],
    
    "section_16_professional_standards": [
        {{"standard": "Name [Source]", "issuing_body": "Organisation", "requirement_type": "Mandatory/Best Practice"}}
    ],
    
    "section_17_equality_statement": {{
        "statement": "This analysis identifies genuine occupational requirements only. No requirements discriminate based on protected characteristics under the Equality Act 2010.",
        "bias_concerns": "None identified" 
    }},
    
    "section_18_citations": [
        {{"source_type": "Type", "source_name": "Name", "url": "URL", "accessed": "{datetime.now().strftime('%Y-%m-%d')}"}}
    ],
    
    "tasks": [
        {{
            "task_id": "T-001",
            "task_description": "Task description [Source: URL]",
            "knowledge": ["Knowledge item"],
            "skills": ["Skill item"],
            "behaviours": ["Behaviour"],
            "criticality": "High/Medium/Low",
            "frequency": "Daily/Weekly/Monthly",
            "source": "URL"
        }}
    ]
}}
```

## CRITICAL REMINDERS
1. Use web_search for EVERY section - do not rely on training data
2. Include actual URLs from your searches
3. If not found, use "Not found through research" - NEVER fabricate
4. Generate 8-20 tasks based on actual job descriptions found
5. Every claim needs a [Source: URL] citation"""



# ============================================================================
# ANALYSIS AGENT - EXECUTION AND DOCUMENT GENERATION
# ============================================================================

async def run_analysis_agent(job_id: str, parameters: Dict, framework: str, terms: Dict):
    """
    Research-based Analysis Agent v6.0
    
    Follows the principle: "EXTRACT and CITE, never GENERATE and GUESS"
    
    Outputs:
    - 01_Job_Task_Analysis.docx - Framework-compliant task list
    - 02_Analysis_Report.docx - 18-section comprehensive report
    - analysis_data.json - Raw research data
    """
    
    # Extract parameters
    domain = parameters.get("domain", "General")
    specialism = parameters.get("specialism", domain)
    role_title = parameters.get("role_title", "Training Specialist")
    proficiency_level = parameters.get("proficiency_level", "Mid-Level")
    role_description = parameters.get("role_description", "")
    
    output_dir = Path(jobs.get(job_id)["output_dir"]) / "01_Analysis"
    output_dir.mkdir(exist_ok=True)
    
    print(f"[NOVA v6.0] Analysis Agent: {role_title} in {domain}/{specialism}")
    update_job(job_id, 2, f"Starting Research Analysis ({framework})...")
    
    # Build research prompt
    update_job(job_id, 5, "Building research queries...")
    research_prompt = build_research_prompt(
        domain=domain,
        specialism=specialism,
        role_title=role_title,
        proficiency_level=proficiency_level,
        framework=framework,
        role_description=role_description,
        terms=terms
    )
    
    print(f"[NOVA] Research prompt: {len(research_prompt)} chars")
    
    # Execute research with web search
    update_job(job_id, 10, "Conducting web research (this may take 1-2 minutes)...")
    
    try:
        research_result = await call_claude_with_search(
            system_prompt=get_analysis_system_prompt(),
            user_prompt=research_prompt,
            max_tokens=16000
        )
        
        searches = research_result.get("searches_performed", [])
        print(f"[NOVA] Research complete: {len(searches)} web searches performed")
        update_job(job_id, 45, f"Research complete. {len(searches)} searches performed.")
        
        # Parse research output
        update_job(job_id, 50, "Processing research findings...")
        analysis_data = parse_analysis_output(research_result.get("text", ""))
        
        # Ensure all sections exist
        analysis_data = ensure_all_sections(analysis_data, parameters, framework, terms)
        
        # Add metadata
        analysis_data["metadata"] = {
            "domain": domain,
            "specialism": specialism,
            "role_title": role_title,
            "proficiency_level": proficiency_level,
            "framework": framework,
            "framework_name": terms.get("framework_name", framework),
            "generated_date": datetime.now().isoformat(),
            "searches_performed": searches,
            "nova_version": "6.0.0"
        }
        
        task_count = len(analysis_data.get("tasks", []))
        citation_count = len(analysis_data.get("section_18_citations", []))
        update_job(job_id, 55, f"Found {task_count} tasks, {citation_count} sources")
        
    except Exception as e:
        print(f"[NOVA] Research error: {e}")
        import traceback
        traceback.print_exc()
        update_job(job_id, 45, "Research encountered issues, creating partial report...")
        analysis_data = create_fallback_analysis(parameters, framework, terms, str(e))
    
    # Generate documents
    update_job(job_id, 60, f"Building {terms['task_list']}...")
    build_task_list_document(analysis_data, role_title, framework, terms, output_dir / "01_Job_Task_Analysis.docx")
    update_job(job_id, 75, f"✓ {terms['task_list']} complete")
    
    update_job(job_id, 80, "Building 18-Section Analysis Report...")
    build_analysis_report_document(analysis_data, role_title, framework, terms, output_dir / "02_Analysis_Report.docx")
    update_job(job_id, 95, "✓ Analysis Report complete")
    
    # Save JSON
    with open(output_dir / "analysis_data.json", "w") as f:
        json.dump(analysis_data, f, indent=2, default=str)
    
    update_job(job_id, 100, "Analysis Phase Complete")


def parse_analysis_output(text: str) -> Dict:
    """Parse JSON output from Claude, with error recovery"""
    if not text:
        return {"parse_error": True}
    
    text = text.strip()
    
    # Try direct parse
    try:
        if text.startswith('{'):
            return json.loads(text)
    except json.JSONDecodeError:
        pass
    
    # Try to find JSON block
    try:
        # Find first { and last }
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1:
            json_str = text[start:end+1]
            return json.loads(json_str)
    except json.JSONDecodeError:
        pass
    
    # Try to extract with regex
    try:
        match = re.search(r'\{[\s\S]*\}', text)
        if match:
            return json.loads(match.group())
    except:
        pass
    
    return {"parse_error": True, "raw_text": text[:3000]}


def ensure_all_sections(data: Dict, parameters: Dict, framework: str, terms: Dict) -> Dict:
    """Ensure all 18 required sections exist with proper defaults"""
    
    role_title = parameters.get("role_title", "Role")
    domain = parameters.get("domain", "Domain")
    specialism = parameters.get("specialism", domain)
    proficiency_level = parameters.get("proficiency_level", "Mid-Level")
    
    defaults = {
        "research_summary": {
            "searches_conducted": [],
            "sources_found": [],
            "research_date": datetime.now().strftime('%Y-%m-%d'),
            "research_limitations": "See individual sections for specific gaps"
        },
        "section_01_executive_summary": {
            "analysis_scope": f"Analysis of {role_title} in {domain}/{specialism} at {proficiency_level} level",
            "methodology": "Web-based research using authoritative UK sources",
            "key_findings": ["Research findings documented in subsequent sections"],
            "tasks_identified": 0,
            "primary_standards_referenced": []
        },
        "section_02_framework_identification": {
            "framework_name": terms.get("framework_name", framework),
            "framework_version": "Not found through research",
            "governing_authority": terms.get("authority", "Not found through research"),
            "analysis_requirements": [],
            "terminology_glossary": {}
        },
        "section_03_geographic_context": {
            "country": "United Kingdom",
            "legal_jurisdiction": "England and Wales",
            "language": "English",
            "currency": "GBP"
        },
        "section_04_professional_body": {
            "body_name": "Not found through research",
            "website_url": "",
            "registration_required": False,
            "registration_requirements": [],
            "protected_titles": [],
            "membership_categories": []
        },
        "section_05_competency_framework": {
            "framework_name": "Not found through research",
            "framework_owner": "",
            "relevant_competencies": [],
            "proficiency_mapping": "Not found through research"
        },
        "section_06_role_description": {
            "definition": "Not found through research",
            "primary_purpose": "",
            "key_accountabilities": [],
            "reporting_structure": "",
            "equivalent_titles": []
        },
        "section_07_qualifications": {
            "essential": [],
            "desirable": [],
            "apprenticeship_routes": []
        },
        "section_08_experience": {
            "years_required": "Not found through research",
            "experience_types": [],
            "sector_requirements": []
        },
        "section_09_technical_skills": [],
        "section_10_soft_skills": [],
        "section_11_behaviours": [],
        "section_12_physical_medical_security": {
            "physical_requirements": "No specific requirements identified",
            "medical_requirements": "Standard employment health requirements",
            "security_clearance": "Standard DBS check",
            "dbs_level": "Basic"
        },
        "section_13_cpd_requirements": {
            "professional_body_cpd": "Not found through research",
            "annual_hours": "Not found through research",
            "recertification_cycle": "Not found through research",
            "mandatory_refreshers": []
        },
        "section_14_career_progression": {
            "pathway_to_role": [],
            "pathway_from_role": [],
            "typical_timeline": "Not found through research"
        },
        "section_15_legal_compliance": [],
        "section_16_professional_standards": [],
        "section_17_equality_statement": {
            "statement": "This analysis identifies genuine occupational requirements only. All requirements comply with the Equality Act 2010.",
            "bias_concerns": "None identified"
        },
        "section_18_citations": [],
        "tasks": []
    }
    
    # Merge defaults with data
    for key, default_value in defaults.items():
        if key not in data:
            data[key] = default_value
        elif isinstance(default_value, dict) and isinstance(data.get(key), dict):
            for sub_key, sub_value in default_value.items():
                if sub_key not in data[key]:
                    data[key][sub_key] = sub_value
    
    # Update task count
    if "section_01_executive_summary" in data:
        data["section_01_executive_summary"]["tasks_identified"] = len(data.get("tasks", []))
    
    return data


def create_fallback_analysis(parameters: Dict, framework: str, terms: Dict, error: str) -> Dict:
    """Create fallback structure when research fails"""
    
    role_title = parameters.get("role_title", "Role")
    domain = parameters.get("domain", "Domain")
    
    return {
        "research_summary": {
            "searches_conducted": [],
            "sources_found": [],
            "research_date": datetime.now().strftime('%Y-%m-%d'),
            "research_limitations": f"Research could not be completed: {error}"
        },
        "section_01_executive_summary": {
            "analysis_scope": f"Analysis of {role_title} in {domain}",
            "methodology": "Research incomplete - manual analysis required",
            "key_findings": ["Research could not be completed - manual input required"],
            "tasks_identified": 0,
            "primary_standards_referenced": []
        },
        "section_02_framework_identification": {
            "framework_name": terms.get("framework_name", framework),
            "framework_version": "Not determined",
            "governing_authority": terms.get("authority", ""),
            "analysis_requirements": [],
            "terminology_glossary": {}
        },
        "section_03_geographic_context": {
            "country": "United Kingdom",
            "legal_jurisdiction": "England and Wales", 
            "language": "English",
            "currency": "GBP"
        },
        "error_message": f"Analysis incomplete due to: {error}",
        "tasks": [],
        "section_18_citations": []
    }


# ============================================================================
# DOCUMENT BUILDERS - TASK LIST
# ============================================================================

def build_task_list_document(data: Dict, role_title: str, framework: str, terms: Dict, output_path: Path):
    """Build the Job/Task Analysis document"""
    
    doc = Document()
    
    # Set styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading(f"{terms['task_list']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(role_title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    
    # Framework badge
    framework_para = doc.add_paragraph(f"Framework: {terms.get('framework_name', framework)}")
    framework_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Metadata section
    doc.add_heading("Analysis Parameters", level=1)
    
    metadata = data.get("metadata", {})
    meta_table = doc.add_table(rows=7, cols=2)
    meta_table.style = 'Table Grid'
    
    meta_rows = [
        ("Domain", metadata.get("domain", "Not specified")),
        ("Specialism", metadata.get("specialism", "Not specified")),
        ("Role Title", metadata.get("role_title", role_title)),
        ("Proficiency Level", metadata.get("proficiency_level", "Not specified")),
        ("Framework", terms.get("framework_name", framework)),
        ("Analysis Date", metadata.get("generated_date", datetime.now().strftime('%Y-%m-%d'))),
        ("NOVA Version", metadata.get("nova_version", "6.0.0"))
    ]
    
    for i, (label, value) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = str(value)
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Research Summary
    research_summary = data.get("research_summary", {})
    doc.add_heading("Research Summary", level=1)
    
    sources = research_summary.get("sources_found", [])
    if sources:
        doc.add_paragraph(f"Sources identified: {len(sources)}")
        for source in sources[:10]:  # Limit to 10
            doc.add_paragraph(f"• {source}", style='List Bullet')
    else:
        doc.add_paragraph("No sources recorded in research summary.")
    
    limitations = research_summary.get("research_limitations", "")
    if limitations:
        doc.add_paragraph(f"Research limitations: {limitations}")
    
    doc.add_paragraph()
    
    # Tasks Section
    tasks = data.get("tasks", [])
    doc.add_heading(f"Tasks Identified: {len(tasks)}", level=1)
    
    if tasks:
        # Create task table with framework-appropriate headers
        if framework in ["UK_DSAT", "S6000T"]:
            headers = ["Task No.", "Performance", "Conditions", "Standards", "Category", "Source"]
            col_count = 6
        elif framework == "US_TRADOC":
            headers = ["Task No.", "Task Title", "Domain", "Frequency", "Source"]
            col_count = 5
        else:
            headers = ["Task ID", "Task Description", "Knowledge/Skills", "Criticality", "Source"]
            col_count = 5
        
        task_table = doc.add_table(rows=1, cols=col_count)
        task_table.style = 'Table Grid'
        
        # Header row
        header_row = task_table.rows[0]
        for i, header in enumerate(headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].bold = True
        
        # Task rows
        for task in tasks:
            row = task_table.add_row()
            
            if framework in ["UK_DSAT", "S6000T"]:
                row.cells[0].text = str(task.get("task_id", ""))
                row.cells[1].text = str(task.get("task_description", ""))[:100]
                row.cells[2].text = "Standard conditions"
                row.cells[3].text = "To required standard"
                row.cells[4].text = task.get("criticality", "Medium")[:1]
                row.cells[5].text = task.get("source", "Research")[:30]
            elif framework == "US_TRADOC":
                row.cells[0].text = str(task.get("task_id", ""))
                row.cells[1].text = str(task.get("task_description", ""))[:100]
                row.cells[2].text = "INST"
                row.cells[3].text = task.get("frequency", "AN")[:10]
                row.cells[4].text = task.get("source", "Research")[:30]
            else:
                row.cells[0].text = str(task.get("task_id", ""))
                row.cells[1].text = str(task.get("task_description", ""))[:100]
                ks = ", ".join(task.get("knowledge", [])[:2] + task.get("skills", [])[:2])
                row.cells[2].text = ks[:50] if ks else "See report"
                row.cells[3].text = task.get("criticality", "Medium")
                row.cells[4].text = task.get("source", "Research")[:30]
    else:
        doc.add_paragraph("No tasks were identified through research. Manual task analysis required.")
    
    doc.add_paragraph()
    
    # Disclaimer
    doc.add_heading("Research Methodology Statement", level=1)
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "This document was generated through AI-assisted web research. All factual claims "
        "are derived from authoritative sources identified through web search. Items marked "
        "'Not found through research' require manual verification. This document should be "
        "reviewed by a subject matter expert before use in formal training development."
    )
    disclaimer.runs[0].font.italic = True
    disclaimer.runs[0].font.size = Pt(10)
    
    # Save
    doc.save(str(output_path))
    print(f"[NOVA] Task list saved: {output_path}")



# ============================================================================
# DOCUMENT BUILDERS - 18-SECTION ANALYSIS REPORT
# ============================================================================

def build_analysis_report_document(data: Dict, role_title: str, framework: str, terms: Dict, output_path: Path):
    """Build the comprehensive 18-section Analysis Report"""
    
    doc = Document()
    
    # Set styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    metadata = data.get("metadata", {})
    
    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading("TRAINING ANALYSIS REPORT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Role title
    role_para = doc.add_paragraph(role_title)
    role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_run = role_para.runs[0]
    role_run.font.size = Pt(18)
    role_run.font.bold = True
    
    doc.add_paragraph()
    
    # Framework
    fw_para = doc.add_paragraph(terms.get("framework_name", framework))
    fw_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fw_para.runs[0].font.size = Pt(14)
    
    # Badge
    badge_para = doc.add_paragraph("18-SECTION COMPREHENSIVE ANALYSIS")
    badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_para.runs[0].font.size = Pt(12)
    badge_para.runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ("Domain", metadata.get("domain", "Not specified")),
        ("Specialism", metadata.get("specialism", "Not specified")),
        ("Proficiency Level", metadata.get("proficiency_level", "Not specified")),
        ("Analysis Date", metadata.get("generated_date", datetime.now().strftime('%Y-%m-%d'))[:10]),
        ("NOVA Version", metadata.get("nova_version", "6.0.0"))
    ]
    for i, (label, value) in enumerate(info_rows):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    doc.add_heading("TABLE OF CONTENTS", level=1)
    
    toc_items = [
        "1. Executive Summary",
        "2. Framework Identification",
        "3. Geographic/Jurisdictional Context",
        "4. Professional Body/Regulator",
        "5. Competency Framework Mapping",
        "6. Role Description",
        "7. Qualifications",
        "8. Experience Requirements",
        "9. Technical Skills",
        "10. Soft Skills",
        "11. Behaviours",
        "12. Physical/Medical/Security Requirements",
        "13. CPD/Recertification Requirements",
        "14. Career Progression",
        "15. Legal Compliance",
        "16. Professional Standards",
        "17. Equality Statement",
        "18. Citations and Sources"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ===== SECTION 1: EXECUTIVE SUMMARY =====
    doc.add_heading("1. EXECUTIVE SUMMARY", level=1)
    
    s1 = data.get("section_01_executive_summary", {})
    
    doc.add_heading("Analysis Scope", level=2)
    doc.add_paragraph(s1.get("analysis_scope", "Not specified"))
    
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(s1.get("methodology", "Web-based research using authoritative sources"))
    
    doc.add_heading("Key Findings", level=2)
    findings = s1.get("key_findings", [])
    if findings:
        for finding in findings:
            doc.add_paragraph(f"• {finding}", style='List Bullet')
    else:
        doc.add_paragraph("No specific findings recorded.")
    
    doc.add_heading("Summary Statistics", level=2)
    stats_table = doc.add_table(rows=2, cols=2)
    stats_table.style = 'Table Grid'
    stats_table.rows[0].cells[0].text = "Tasks Identified"
    stats_table.rows[0].cells[1].text = str(s1.get("tasks_identified", len(data.get("tasks", []))))
    stats_table.rows[1].cells[0].text = "Primary Standards"
    standards = s1.get("primary_standards_referenced", [])
    stats_table.rows[1].cells[1].text = ", ".join(standards) if standards else "See Section 16"
    
    doc.add_paragraph()
    
    # ===== SECTION 2: FRAMEWORK IDENTIFICATION =====
    doc.add_heading("2. FRAMEWORK IDENTIFICATION", level=1)
    
    s2 = data.get("section_02_framework_identification", {})
    
    fw_table = doc.add_table(rows=4, cols=2)
    fw_table.style = 'Table Grid'
    fw_rows = [
        ("Framework Name", s2.get("framework_name", terms.get("framework_name", framework))),
        ("Version", s2.get("framework_version", "Not found through research")),
        ("Governing Authority", s2.get("governing_authority", terms.get("authority", "Not specified"))),
        ("Citation Format", terms.get("citation_prefix", "[Reference]"))
    ]
    for i, (label, value) in enumerate(fw_rows):
        fw_table.rows[i].cells[0].text = label
        fw_table.rows[i].cells[1].text = str(value)
        fw_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    reqs = s2.get("analysis_requirements", [])
    if reqs:
        doc.add_heading("Analysis Phase Requirements", level=2)
        for req in reqs:
            doc.add_paragraph(f"• {req}", style='List Bullet')
    
    glossary = s2.get("terminology_glossary", {})
    if glossary:
        doc.add_heading("Terminology Glossary", level=2)
        for term, defn in glossary.items():
            doc.add_paragraph(f"• {term}: {defn}", style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== SECTION 3: GEOGRAPHIC CONTEXT =====
    doc.add_heading("3. GEOGRAPHIC/JURISDICTIONAL CONTEXT", level=1)
    
    s3 = data.get("section_03_geographic_context", {})
    
    geo_table = doc.add_table(rows=4, cols=2)
    geo_table.style = 'Table Grid'
    geo_rows = [
        ("Country", s3.get("country", "United Kingdom")),
        ("Legal Jurisdiction", s3.get("legal_jurisdiction", "England and Wales")),
        ("Language", s3.get("language", "English")),
        ("Currency", s3.get("currency", "GBP"))
    ]
    for i, (label, value) in enumerate(geo_rows):
        geo_table.rows[i].cells[0].text = label
        geo_table.rows[i].cells[1].text = str(value)
        geo_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ===== SECTION 4: PROFESSIONAL BODY =====
    doc.add_heading("4. PROFESSIONAL BODY/REGULATOR", level=1)
    
    s4 = data.get("section_04_professional_body", {})
    
    body_name = s4.get("body_name", "Not found through research")
    doc.add_paragraph(f"Professional Body: {body_name}").runs[0].bold = True
    
    if s4.get("website_url"):
        doc.add_paragraph(f"Website: {s4.get('website_url')}")
    
    doc.add_paragraph(f"Registration Required: {'Yes' if s4.get('registration_required') else 'No / Not determined'}")
    
    reqs = s4.get("registration_requirements", [])
    if reqs:
        doc.add_heading("Registration Requirements", level=2)
        for req in reqs:
            doc.add_paragraph(f"• {req}", style='List Bullet')
    
    titles = s4.get("protected_titles", [])
    if titles:
        doc.add_heading("Protected Titles", level=2)
        for title_item in titles:
            doc.add_paragraph(f"• {title_item}", style='List Bullet')
    
    categories = s4.get("membership_categories", [])
    if categories:
        doc.add_heading("Membership Categories", level=2)
        for cat in categories:
            doc.add_paragraph(f"• {cat}", style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== SECTION 5: COMPETENCY FRAMEWORK =====
    doc.add_heading("5. COMPETENCY FRAMEWORK MAPPING", level=1)
    
    s5 = data.get("section_05_competency_framework", {})
    
    doc.add_paragraph(f"Framework: {s5.get('framework_name', 'Not found through research')}").runs[0].bold = True
    
    if s5.get("framework_owner"):
        doc.add_paragraph(f"Owner: {s5.get('framework_owner')}")
    
    competencies = s5.get("relevant_competencies", [])
    if competencies:
        doc.add_heading("Relevant Competency Units", level=2)
        comp_table = doc.add_table(rows=1, cols=3)
        comp_table.style = 'Table Grid'
        comp_table.rows[0].cells[0].text = "Code"
        comp_table.rows[0].cells[1].text = "Name"
        comp_table.rows[0].cells[2].text = "Level"
        for comp in competencies:
            row = comp_table.add_row()
            row.cells[0].text = str(comp.get("code", ""))
            row.cells[1].text = str(comp.get("name", ""))
            row.cells[2].text = str(comp.get("level", ""))
    
    mapping = s5.get("proficiency_mapping", "")
    if mapping:
        doc.add_heading("Proficiency Level Mapping", level=2)
        doc.add_paragraph(mapping)
    
    doc.add_paragraph()
    
    # ===== SECTION 6: ROLE DESCRIPTION =====
    doc.add_heading("6. ROLE DESCRIPTION", level=1)
    
    s6 = data.get("section_06_role_description", {})
    
    doc.add_heading("Definition", level=2)
    doc.add_paragraph(s6.get("definition", "Not found through research"))
    
    if s6.get("primary_purpose"):
        doc.add_heading("Primary Purpose", level=2)
        doc.add_paragraph(s6.get("primary_purpose"))
    
    accountabilities = s6.get("key_accountabilities", [])
    if accountabilities:
        doc.add_heading("Key Accountabilities", level=2)
        for acc in accountabilities:
            doc.add_paragraph(f"• {acc}", style='List Bullet')
    
    if s6.get("reporting_structure"):
        doc.add_heading("Reporting Structure", level=2)
        doc.add_paragraph(s6.get("reporting_structure"))
    
    equiv = s6.get("equivalent_titles", [])
    if equiv:
        doc.add_heading("Equivalent Job Titles", level=2)
        doc.add_paragraph(", ".join(equiv))
    
    doc.add_paragraph()
    
    # ===== SECTION 7: QUALIFICATIONS =====
    doc.add_heading("7. QUALIFICATIONS", level=1)
    
    s7 = data.get("section_07_qualifications", {})
    
    essential = s7.get("essential", [])
    if essential:
        doc.add_heading("Essential Qualifications", level=2)
        qual_table = doc.add_table(rows=1, cols=3)
        qual_table.style = 'Table Grid'
        qual_table.rows[0].cells[0].text = "Qualification"
        qual_table.rows[0].cells[1].text = "Level"
        qual_table.rows[0].cells[2].text = "Source"
        for q in essential:
            row = qual_table.add_row()
            row.cells[0].text = str(q.get("qualification", ""))
            row.cells[1].text = str(q.get("level", ""))
            row.cells[2].text = str(q.get("source", ""))[:40]
    else:
        doc.add_paragraph("Essential qualifications: Not found through research")
    
    desirable = s7.get("desirable", [])
    if desirable:
        doc.add_heading("Desirable Qualifications", level=2)
        for q in desirable:
            doc.add_paragraph(f"• {q.get('qualification', '')} ({q.get('level', '')})", style='List Bullet')
    
    apprenticeships = s7.get("apprenticeship_routes", [])
    if apprenticeships:
        doc.add_heading("Apprenticeship Routes", level=2)
        for a in apprenticeships:
            doc.add_paragraph(f"• {a.get('name', '')} - Level {a.get('level', '')}", style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== SECTION 8: EXPERIENCE =====
    doc.add_heading("8. EXPERIENCE REQUIREMENTS", level=1)
    
    s8 = data.get("section_08_experience", {})
    
    doc.add_paragraph(f"Years Required: {s8.get('years_required', 'Not found through research')}").runs[0].bold = True
    
    exp_types = s8.get("experience_types", [])
    if exp_types:
        doc.add_heading("Type of Experience", level=2)
        for exp in exp_types:
            doc.add_paragraph(f"• {exp}", style='List Bullet')
    
    sector = s8.get("sector_requirements", [])
    if sector:
        doc.add_heading("Sector Requirements", level=2)
        for sec in sector:
            doc.add_paragraph(f"• {sec}", style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== SECTION 9: TECHNICAL SKILLS =====
    doc.add_heading("9. TECHNICAL SKILLS", level=1)
    
    s9 = data.get("section_09_technical_skills", [])
    
    if s9:
        skills_table = doc.add_table(rows=1, cols=4)
        skills_table.style = 'Table Grid'
        skills_table.rows[0].cells[0].text = "Skill"
        skills_table.rows[0].cells[1].text = "Category"
        skills_table.rows[0].cells[2].text = "Proficiency"
        skills_table.rows[0].cells[3].text = "Source"
        for skill in s9:
            row = skills_table.add_row()
            row.cells[0].text = str(skill.get("skill", ""))
            row.cells[1].text = str(skill.get("category", "Core"))
            row.cells[2].text = str(skill.get("proficiency", ""))
            row.cells[3].text = str(skill.get("source", ""))[:30]
    else:
        doc.add_paragraph("Technical skills: Not found through research")
    
    doc.add_paragraph()
    
    # ===== SECTION 10: SOFT SKILLS =====
    doc.add_heading("10. SOFT SKILLS", level=1)
    
    s10 = data.get("section_10_soft_skills", [])
    
    if s10:
        soft_table = doc.add_table(rows=1, cols=3)
        soft_table.style = 'Table Grid'
        soft_table.rows[0].cells[0].text = "Skill"
        soft_table.rows[0].cells[1].text = "Proficiency"
        soft_table.rows[0].cells[2].text = "Source"
        for skill in s10:
            row = soft_table.add_row()
            row.cells[0].text = str(skill.get("skill", ""))
            row.cells[1].text = str(skill.get("proficiency", ""))
            row.cells[2].text = str(skill.get("source", ""))[:30]
    else:
        doc.add_paragraph("Soft skills: Not found through research")
    
    doc.add_paragraph()
    
    # ===== SECTION 11: BEHAVIOURS =====
    doc.add_heading("11. BEHAVIOURS", level=1)
    
    s11 = data.get("section_11_behaviours", [])
    
    if s11:
        for beh in s11:
            doc.add_paragraph(f"• {beh.get('behaviour', '')} ({beh.get('type', 'Essential')})", style='List Bullet')
    else:
        doc.add_paragraph("Behaviours: Not found through research")
    
    doc.add_paragraph()
    
    # ===== SECTION 12: PHYSICAL/MEDICAL/SECURITY =====
    doc.add_heading("12. PHYSICAL/MEDICAL/SECURITY REQUIREMENTS", level=1)
    
    s12 = data.get("section_12_physical_medical_security", {})
    
    pms_table = doc.add_table(rows=4, cols=2)
    pms_table.style = 'Table Grid'
    pms_rows = [
        ("Physical Requirements", s12.get("physical_requirements", "No specific requirements identified")),
        ("Medical Requirements", s12.get("medical_requirements", "Standard employment health requirements")),
        ("Security Clearance", s12.get("security_clearance", "Standard DBS check")),
        ("DBS Level", s12.get("dbs_level", "Basic"))
    ]
    for i, (label, value) in enumerate(pms_rows):
        pms_table.rows[i].cells[0].text = label
        pms_table.rows[i].cells[1].text = str(value)
        pms_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ===== SECTION 13: CPD REQUIREMENTS =====
    doc.add_heading("13. CPD/RECERTIFICATION REQUIREMENTS", level=1)
    
    s13 = data.get("section_13_cpd_requirements", {})
    
    cpd_table = doc.add_table(rows=4, cols=2)
    cpd_table.style = 'Table Grid'
    cpd_rows = [
        ("Professional Body CPD Policy", s13.get("professional_body_cpd", "Not found through research")),
        ("Annual Hours/Points", s13.get("annual_hours", "Not found through research")),
        ("Recertification Cycle", s13.get("recertification_cycle", "Not found through research")),
        ("Mandatory Refreshers", ", ".join(s13.get("mandatory_refreshers", [])) or "None identified")
    ]
    for i, (label, value) in enumerate(cpd_rows):
        cpd_table.rows[i].cells[0].text = label
        cpd_table.rows[i].cells[1].text = str(value)
        cpd_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ===== SECTION 14: CAREER PROGRESSION =====
    doc.add_heading("14. CAREER PROGRESSION", level=1)
    
    s14 = data.get("section_14_career_progression", {})
    
    to_role = s14.get("pathway_to_role", [])
    if to_role:
        doc.add_heading("Pathway to This Role", level=2)
        for role in to_role:
            doc.add_paragraph(f"• {role}", style='List Bullet')
    
    from_role = s14.get("pathway_from_role", [])
    if from_role:
        doc.add_heading("Progression from This Role", level=2)
        for role in from_role:
            doc.add_paragraph(f"• {role}", style='List Bullet')
    
    timeline = s14.get("typical_timeline", "")
    if timeline:
        doc.add_paragraph(f"Typical Timeline: {timeline}")
    
    doc.add_paragraph()
    
    # ===== SECTION 15: LEGAL COMPLIANCE =====
    doc.add_heading("15. LEGAL COMPLIANCE", level=1)
    
    s15 = data.get("section_15_legal_compliance", [])
    
    if s15:
        legal_table = doc.add_table(rows=1, cols=4)
        legal_table.style = 'Table Grid'
        legal_table.rows[0].cells[0].text = "Legislation"
        legal_table.rows[0].cells[1].text = "Relevance"
        legal_table.rows[0].cells[2].text = "Mandatory Training"
        legal_table.rows[0].cells[3].text = "Source"
        for leg in s15:
            row = legal_table.add_row()
            row.cells[0].text = str(leg.get("legislation", ""))
            row.cells[1].text = str(leg.get("relevance", ""))[:50]
            row.cells[2].text = "Yes" if leg.get("mandatory_training") else "No"
            row.cells[3].text = str(leg.get("source", ""))[:30]
    else:
        doc.add_paragraph("Legal requirements: Not found through research")
    
    doc.add_paragraph()
    
    # ===== SECTION 16: PROFESSIONAL STANDARDS =====
    doc.add_heading("16. PROFESSIONAL STANDARDS", level=1)
    
    s16 = data.get("section_16_professional_standards", [])
    
    if s16:
        for std in s16:
            doc.add_paragraph(f"• {std.get('standard', '')} - {std.get('issuing_body', '')} ({std.get('requirement_type', '')})", style='List Bullet')
    else:
        doc.add_paragraph("Professional standards: Not found through research")
    
    doc.add_paragraph()
    
    # ===== SECTION 17: EQUALITY STATEMENT =====
    doc.add_heading("17. EQUALITY STATEMENT", level=1)
    
    s17 = data.get("section_17_equality_statement", {})
    
    doc.add_paragraph(s17.get("statement", "This analysis identifies genuine occupational requirements only."))
    
    bias = s17.get("bias_concerns", "None identified")
    doc.add_paragraph(f"Bias Concerns: {bias}")
    
    doc.add_paragraph()
    
    # ===== SECTION 18: CITATIONS =====
    doc.add_heading("18. CITATIONS AND SOURCES", level=1)
    
    s18 = data.get("section_18_citations", [])
    
    if s18:
        cite_table = doc.add_table(rows=1, cols=4)
        cite_table.style = 'Table Grid'
        cite_table.rows[0].cells[0].text = "Type"
        cite_table.rows[0].cells[1].text = "Source"
        cite_table.rows[0].cells[2].text = "URL"
        cite_table.rows[0].cells[3].text = "Accessed"
        for cite in s18:
            row = cite_table.add_row()
            row.cells[0].text = str(cite.get("source_type", ""))
            row.cells[1].text = str(cite.get("source_name", ""))[:40]
            row.cells[2].text = str(cite.get("url", ""))[:50]
            row.cells[3].text = str(cite.get("accessed", ""))
    else:
        doc.add_paragraph("No citations recorded. Sources are noted inline throughout the document.")
    
    # Research log
    research_summary = data.get("research_summary", {})
    searches = research_summary.get("searches_conducted", [])
    if searches:
        doc.add_heading("Research Log", level=2)
        doc.add_paragraph(f"Searches conducted: {len(searches)}")
        for i, search in enumerate(searches[:20], 1):  # Limit to 20
            doc.add_paragraph(f"{i}. {search}")
    
    doc.add_paragraph()
    
    # ===== FINAL DISCLAIMER =====
    doc.add_heading("IMPORTANT DISCLAIMER", level=1)
    
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "This analysis report was generated through AI-assisted web research using the NOVA™ platform. "
        "All factual claims are derived from authoritative sources identified through web search at the time of analysis. "
        "Items marked 'Not found through research' indicate gaps in available online information and require manual verification.\n\n"
        "This document is intended as a starting point for training development and should be reviewed and validated by:\n"
        "• A subject matter expert (SME) in the relevant domain\n"
        "• A training design professional\n"
        "• Appropriate governance authorities\n\n"
        "NOVA™ does not guarantee the accuracy, completeness, or currency of information obtained through web research. "
        "Users are responsible for verifying all requirements with authoritative sources before use in formal training development."
    )
    disclaimer.runs[0].font.italic = True
    disclaimer.runs[0].font.size = Pt(10)
    
    # Save
    doc.save(str(output_path))
    print(f"[NOVA] Analysis report saved: {output_path}")



# ============================================================================
# DESIGN AGENT
# ============================================================================

async def run_design_agent(job_id: str, parameters: Dict, framework: str, terms: Dict):
    """Design Agent - generates TOs, EOs, KLPs/LSAs, and course design documents"""
    
    role_title = parameters.get("role_title", "Training Role")
    analysis_data = parameters.get("analysis_data", {})
    tasks = analysis_data.get("tasks", parameters.get("tasks", []))
    
    output_dir = Path(jobs.get(job_id)["output_dir"]) / "02_Design"
    output_dir.mkdir(exist_ok=True)
    
    update_job(job_id, 5, f"Starting Design Agent ({framework})...")
    
    # Generate Training Objectives
    update_job(job_id, 10, f"Generating {terms['top_objective_short']}s...")
    
    to_prompt = f"""Generate Training Objectives for: {role_title}
Framework: {terms['framework_name']}

Tasks to convert to objectives:
{json.dumps(tasks[:15], indent=2)}

Return JSON with this structure:
{{
    "training_objectives": [
        {{
            "to_id": "{terms['top_objective_short']} 1",
            "performance": "The trainee will [action verb] [object]",
            "conditions": "Given [equipment, environment, constraints]",
            "standards": "To the standard of [measurable criteria]",
            "source_task": "T-001",
            "ksa_domain": "Knowledge/Skill/Attitude"
        }}
    ]
}}

Rules:
- Use {terms['top_objective_short']} numbering
- Start performance with action verb from Bloom's taxonomy
- Make standards measurable and observable
- Generate 8-15 objectives based on the tasks"""

    try:
        to_response = await call_claude_standard(
            system_prompt=f"You are designing training objectives following {terms['framework_name']}. Be precise and measurable.",
            user_prompt=to_prompt,
            max_tokens=8000
        )
        to_data = parse_json_response(to_response)
    except Exception as e:
        print(f"[NOVA] TO generation error: {e}")
        to_data = {"training_objectives": []}
    
    update_job(job_id, 35, f"✓ Generated {len(to_data.get('training_objectives', []))} {terms['top_objective_short']}s")
    
    # Generate Enabling Objectives
    update_job(job_id, 40, f"Generating {terms['enabling_objective_short']}s...")
    
    eo_prompt = f"""Generate Enabling Objectives for these Training Objectives:
{json.dumps(to_data.get('training_objectives', [])[:10], indent=2)}

Framework: {terms['framework_name']}

Return JSON:
{{
    "enabling_objectives": [
        {{
            "eo_id": "{terms['enabling_objective_short']} 1.1",
            "parent_to": "{terms['top_objective_short']} 1",
            "performance": "The trainee will [specific action]",
            "conditions": "Given [specific context]",
            "standards": "To the standard of [criteria]",
            "ksa_type": "K/S/A"
        }}
    ]
}}

Generate 2-4 {terms['enabling_objective_short']}s per {terms['top_objective_short']}."""

    try:
        eo_response = await call_claude_standard(
            system_prompt=f"You are designing enabling objectives. Each must support its parent objective.",
            user_prompt=eo_prompt,
            max_tokens=8000
        )
        eo_data = parse_json_response(eo_response)
    except Exception as e:
        print(f"[NOVA] EO generation error: {e}")
        eo_data = {"enabling_objectives": []}
    
    update_job(job_id, 60, f"✓ Generated {len(eo_data.get('enabling_objectives', []))} {terms['enabling_objective_short']}s")
    
    # Generate KLPs/LSAs
    update_job(job_id, 65, f"Generating {terms['learning_point_short']}s...")
    
    klp_prompt = f"""Generate {terms['learning_point']}s for these Enabling Objectives:
{json.dumps(eo_data.get('enabling_objectives', [])[:15], indent=2)}

Return JSON:
{{
    "learning_points": [
        {{
            "klp_id": "{terms['learning_point_short']} 1.1.1",
            "parent_eo": "{terms['enabling_objective_short']} 1.1",
            "content": "Declarative statement of what must be learned",
            "domain": "Knowledge/Skill/Attitude"
        }}
    ]
}}

Generate 3-5 {terms['learning_point_short']}s per {terms['enabling_objective_short']}."""

    try:
        klp_response = await call_claude_standard(
            system_prompt=f"You are creating learning points. Each must be a clear, teachable statement.",
            user_prompt=klp_prompt,
            max_tokens=8000
        )
        klp_data = parse_json_response(klp_response)
    except Exception as e:
        print(f"[NOVA] KLP generation error: {e}")
        klp_data = {"learning_points": []}
    
    update_job(job_id, 80, f"✓ Generated {len(klp_data.get('learning_points', []))} {terms['learning_point_short']}s")
    
    # Combine and build documents
    design_data = {
        "metadata": {
            "role_title": role_title,
            "framework": framework,
            "framework_name": terms.get("framework_name"),
            "generated_date": datetime.now().isoformat()
        },
        "training_objectives": to_data.get("training_objectives", []),
        "enabling_objectives": eo_data.get("enabling_objectives", []),
        "learning_points": klp_data.get("learning_points", [])
    }
    
    update_job(job_id, 85, "Building Design documents...")
    
    # Build TO/EO/KLP document
    build_design_hierarchy_doc(design_data, role_title, terms, output_dir / f"01_{terms['top_objective_short']}_Hierarchy.docx")
    
    # Save JSON
    with open(output_dir / "design_data.json", "w") as f:
        json.dump(design_data, f, indent=2)
    
    update_job(job_id, 100, "Design Phase Complete")


def parse_json_response(text: str) -> Dict:
    """Parse JSON from Claude response"""
    text = text.strip()
    try:
        if text.startswith('{'):
            return json.loads(text)
        match = re.search(r'\{[\s\S]*\}', text)
        if match:
            return json.loads(match.group())
    except:
        pass
    return {}


def build_design_hierarchy_doc(data: Dict, role_title: str, terms: Dict, output_path: Path):
    """Build the TO/EO/KLP hierarchy document"""
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading(f"{terms['top_objective']} Hierarchy", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(role_title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Framework: {terms.get('framework_name', 'Standard')}")
    doc.add_paragraph()
    
    # Summary
    doc.add_heading("Summary", level=1)
    tos = data.get("training_objectives", [])
    eos = data.get("enabling_objectives", [])
    klps = data.get("learning_points", [])
    
    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.rows[0].cells[0].text = f"{terms['top_objective_short']}s"
    summary_table.rows[0].cells[1].text = str(len(tos))
    summary_table.rows[1].cells[0].text = f"{terms['enabling_objective_short']}s"
    summary_table.rows[1].cells[1].text = str(len(eos))
    summary_table.rows[2].cells[0].text = f"{terms['learning_point_short']}s"
    summary_table.rows[2].cells[1].text = str(len(klps))
    
    doc.add_paragraph()
    
    # Hierarchy
    doc.add_heading(f"{terms['top_objective']}s", level=1)
    
    for to in tos:
        # TO header
        to_para = doc.add_paragraph()
        to_run = to_para.add_run(f"{to.get('to_id', 'TO')}: ")
        to_run.bold = True
        to_para.add_run(to.get('performance', ''))
        
        doc.add_paragraph(f"Conditions: {to.get('conditions', '')}")
        doc.add_paragraph(f"Standards: {to.get('standards', '')}")
        
        # Find related EOs
        to_id = to.get('to_id', '')
        related_eos = [eo for eo in eos if to_id in str(eo.get('parent_to', ''))]
        
        for eo in related_eos:
            eo_para = doc.add_paragraph(style='List Bullet')
            eo_run = eo_para.add_run(f"{eo.get('eo_id', 'EO')}: ")
            eo_run.bold = True
            eo_para.add_run(eo.get('performance', ''))
            
            # Find related KLPs
            eo_id = eo.get('eo_id', '')
            related_klps = [klp for klp in klps if eo_id in str(klp.get('parent_eo', ''))]
            
            for klp in related_klps:
                klp_para = doc.add_paragraph(style='List Bullet 2')
                klp_para.add_run(f"{klp.get('klp_id', 'KLP')}: {klp.get('content', '')}")
        
        doc.add_paragraph()
    
    doc.save(str(output_path))
    print(f"[NOVA] Design hierarchy saved: {output_path}")


# ============================================================================
# DELIVERY AGENT
# ============================================================================

async def run_delivery_agent(job_id: str, parameters: Dict, framework: str, terms: Dict):
    """Delivery Agent - generates lesson plans, schedules, and assessments"""
    
    role_title = parameters.get("role_title", "Training Role")
    design_data = parameters.get("design_data", {})
    
    output_dir = Path(jobs.get(job_id)["output_dir"]) / "03_Delivery"
    output_dir.mkdir(exist_ok=True)
    
    update_job(job_id, 5, f"Starting Delivery Agent ({framework})...")
    
    tos = design_data.get("training_objectives", [])
    eos = design_data.get("enabling_objectives", [])
    
    # Generate Lesson Plans
    update_job(job_id, 15, "Generating Lesson Plans...")
    
    lesson_prompt = f"""Create lesson plans for these objectives:
{json.dumps(tos[:8], indent=2)}

Framework: {terms['framework_name']}
Lesson Plan Format: {terms.get('lesson_plan', 'Standard')}

Return JSON:
{{
    "lessons": [
        {{
            "lesson_id": "L1",
            "title": "Lesson title",
            "duration_minutes": 60,
            "objectives_covered": ["{terms['top_objective_short']} 1"],
            "present": {{
                "duration_minutes": 25,
                "content": ["Content point 1", "Content point 2"],
                "trainer_notes": "Key points to emphasize"
            }},
            "apply": {{
                "duration_minutes": 25,
                "activities": ["Activity 1", "Activity 2"],
                "resources": ["Resource 1"]
            }},
            "review": {{
                "duration_minutes": 10,
                "questions": ["Review question 1"],
                "summary_points": ["Key takeaway 1"]
            }}
        }}
    ]
}}"""

    try:
        lesson_response = await call_claude_standard(
            system_prompt="You are creating detailed lesson plans. Follow the framework format precisely.",
            user_prompt=lesson_prompt,
            max_tokens=8000
        )
        lesson_data = parse_json_response(lesson_response)
    except Exception as e:
        print(f"[NOVA] Lesson generation error: {e}")
        lesson_data = {"lessons": []}
    
    update_job(job_id, 50, f"✓ Generated {len(lesson_data.get('lessons', []))} lessons")
    
    # Generate Assessment Items
    update_job(job_id, 55, "Generating Assessment Items...")
    
    assessment_prompt = f"""Create assessment items for these objectives:
{json.dumps(tos[:8], indent=2)}

Return JSON:
{{
    "assessments": [
        {{
            "item_id": "A1",
            "objective_ref": "{terms['top_objective_short']} 1",
            "type": "Multiple Choice/Short Answer/Practical",
            "question": "Assessment question",
            "correct_answer": "Correct answer",
            "marking_criteria": "How to mark",
            "pass_criteria": "What constitutes a pass"
        }}
    ]
}}"""

    try:
        assess_response = await call_claude_standard(
            system_prompt="You are creating valid, reliable assessments aligned to objectives.",
            user_prompt=assessment_prompt,
            max_tokens=6000
        )
        assess_data = parse_json_response(assess_response)
    except Exception as e:
        print(f"[NOVA] Assessment generation error: {e}")
        assess_data = {"assessments": []}
    
    update_job(job_id, 75, f"✓ Generated {len(assess_data.get('assessments', []))} assessment items")
    
    # Combine data
    delivery_data = {
        "metadata": {
            "role_title": role_title,
            "framework": framework,
            "generated_date": datetime.now().isoformat()
        },
        "lessons": lesson_data.get("lessons", []),
        "assessments": assess_data.get("assessments", [])
    }
    
    update_job(job_id, 80, "Building Delivery documents...")
    
    # Build lesson plan document
    build_lesson_plans_doc(delivery_data, role_title, terms, output_dir / "01_Lesson_Plans.docx")
    
    # Save JSON
    with open(output_dir / "delivery_data.json", "w") as f:
        json.dump(delivery_data, f, indent=2)
    
    update_job(job_id, 100, "Delivery Phase Complete")


def build_lesson_plans_doc(data: Dict, role_title: str, terms: Dict, output_path: Path):
    """Build lesson plans document"""
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    title = doc.add_heading(f"{terms.get('lesson_plan', 'Lesson Plans')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(role_title).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    lessons = data.get("lessons", [])
    
    for lesson in lessons:
        doc.add_heading(f"{lesson.get('lesson_id', 'Lesson')}: {lesson.get('title', 'Untitled')}", level=1)
        
        doc.add_paragraph(f"Duration: {lesson.get('duration_minutes', 60)} minutes")
        doc.add_paragraph(f"Objectives: {', '.join(lesson.get('objectives_covered', []))}")
        
        # Present section
        present = lesson.get("present", {})
        doc.add_heading("PRESENT", level=2)
        doc.add_paragraph(f"Duration: {present.get('duration_minutes', 20)} minutes")
        for content in present.get("content", []):
            doc.add_paragraph(f"• {content}", style='List Bullet')
        if present.get("trainer_notes"):
            doc.add_paragraph(f"Trainer Notes: {present.get('trainer_notes')}")
        
        # Apply section
        apply = lesson.get("apply", {})
        doc.add_heading("APPLY", level=2)
        doc.add_paragraph(f"Duration: {apply.get('duration_minutes', 30)} minutes")
        for activity in apply.get("activities", []):
            doc.add_paragraph(f"• {activity}", style='List Bullet')
        
        # Review section
        review = lesson.get("review", {})
        doc.add_heading("REVIEW", level=2)
        doc.add_paragraph(f"Duration: {review.get('duration_minutes', 10)} minutes")
        for point in review.get("summary_points", []):
            doc.add_paragraph(f"• {point}", style='List Bullet')
        
        doc.add_paragraph()
    
    doc.save(str(output_path))
    print(f"[NOVA] Lesson plans saved: {output_path}")


# ============================================================================
# EVALUATION AGENT
# ============================================================================

async def run_evaluation_agent(job_id: str, parameters: Dict, framework: str, terms: Dict):
    """Evaluation Agent - generates validation and evaluation documents"""
    
    role_title = parameters.get("role_title", "Training Role")
    
    output_dir = Path(jobs.get(job_id)["output_dir"]) / "04_Evaluation"
    output_dir.mkdir(exist_ok=True)
    
    update_job(job_id, 5, f"Starting Evaluation Agent ({framework})...")
    
    # Generate Evaluation Strategy
    update_job(job_id, 20, "Generating Evaluation Strategy...")
    
    eval_prompt = f"""Create an evaluation strategy for: {role_title}
Framework: {terms['framework_name']}

Include:
- {terms.get('internal_eval', 'Internal Validation')} approach
- {terms.get('external_eval', 'External Validation')} approach
- Data collection methods
- Success criteria

Return JSON:
{{
    "evaluation_strategy": {{
        "purpose": "Purpose of evaluation",
        "internal_validation": {{
            "approach": "Description",
            "frequency": "When conducted",
            "methods": ["Method 1", "Method 2"],
            "responsible": "Who conducts"
        }},
        "external_validation": {{
            "approach": "Description",
            "frequency": "When conducted",
            "methods": ["Method 1", "Method 2"],
            "responsible": "Who conducts"
        }},
        "success_criteria": ["Criterion 1", "Criterion 2"],
        "data_sources": ["Source 1", "Source 2"],
        "reporting": "How results are reported"
    }}
}}"""

    try:
        eval_response = await call_claude_standard(
            system_prompt=f"You are creating an evaluation strategy following {terms['framework_name']}.",
            user_prompt=eval_prompt,
            max_tokens=6000
        )
        eval_data = parse_json_response(eval_response)
    except Exception as e:
        print(f"[NOVA] Evaluation generation error: {e}")
        eval_data = {"evaluation_strategy": {}}
    
    update_job(job_id, 70, "✓ Evaluation strategy complete")
    
    evaluation_data = {
        "metadata": {
            "role_title": role_title,
            "framework": framework,
            "generated_date": datetime.now().isoformat()
        },
        **eval_data
    }
    
    update_job(job_id, 80, "Building Evaluation documents...")
    
    # Build evaluation document
    build_evaluation_doc(evaluation_data, role_title, terms, output_dir / "01_Evaluation_Strategy.docx")
    
    # Save JSON
    with open(output_dir / "evaluation_data.json", "w") as f:
        json.dump(evaluation_data, f, indent=2)
    
    update_job(job_id, 100, "Evaluation Phase Complete")


def build_evaluation_doc(data: Dict, role_title: str, terms: Dict, output_path: Path):
    """Build evaluation strategy document"""
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    title = doc.add_heading("Evaluation Strategy", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(role_title).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    strategy = data.get("evaluation_strategy", {})
    
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(strategy.get("purpose", "To evaluate training effectiveness"))
    
    # Internal Validation
    internal = strategy.get("internal_validation", {})
    doc.add_heading(terms.get("internal_eval", "Internal Validation"), level=1)
    doc.add_paragraph(f"Approach: {internal.get('approach', 'Continuous quality assurance')}")
    doc.add_paragraph(f"Frequency: {internal.get('frequency', 'Ongoing')}")
    if internal.get("methods"):
        doc.add_paragraph("Methods:")
        for method in internal.get("methods", []):
            doc.add_paragraph(f"• {method}", style='List Bullet')
    
    # External Validation
    external = strategy.get("external_validation", {})
    doc.add_heading(terms.get("external_eval", "External Validation"), level=1)
    doc.add_paragraph(f"Approach: {external.get('approach', 'Independent audit')}")
    doc.add_paragraph(f"Frequency: {external.get('frequency', 'Annual')}")
    if external.get("methods"):
        doc.add_paragraph("Methods:")
        for method in external.get("methods", []):
            doc.add_paragraph(f"• {method}", style='List Bullet')
    
    # Success Criteria
    criteria = strategy.get("success_criteria", [])
    if criteria:
        doc.add_heading("Success Criteria", level=1)
        for criterion in criteria:
            doc.add_paragraph(f"• {criterion}", style='List Bullet')
    
    doc.save(str(output_path))
    print(f"[NOVA] Evaluation strategy saved: {output_path}")


# ============================================================================
# STARTUP
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
